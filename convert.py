import re
import json
import os
import sys

# 嘗試匯入必要的庫，如果沒有安裝會提示
try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pyphen
except ImportError:
    pyphen = None

class PetVocabProcessor:
    def __init__(self):
        # 初始化音節拆解工具
        if pyphen:
            self.dic = pyphen.Pyphen(lang='en')
        else:
            self.dic = None
            print("警告: 未安裝 pyphen，將無法自動拆解音節。請執行 pip install pyphen")

    def get_syllables(self, word: str) -> list:
        """
        將單字拆解為音節列表
        例如: 'ability' -> ['a', 'bil', 'i', 'ty']
        """
        if not self.dic or not word:
            return [word] # 如果沒安裝工具，直接回傳原字
        
        # 移除非字母字符 (避免標點符號影響)
        clean_word = re.sub(r'[^a-zA-Z]', '', word)
        if not clean_word: return [word]
        
        hyphenated = self.dic.inserted(clean_word)
        return hyphenated.split('-')

    def clean_word_text(self, text):
        """
        清理單字文字：去除開頭的數字、點、空白
        例如: "1. ability " -> "ability"
        """
        # 去除括號內的詞性 (v.) (n.)
        text = re.sub(r'\(.*?\)', '', text)
        # 去除開頭的數字和點
        text = re.sub(r'^[\d\.]+\s*', '', text)
        return text.strip()

    def parse_docx(self, filename):
        """
        讀取 Word (.docx) 檔案 - 強力相容版
        """
        # 1. 檢查套件
        if not Document:
            print("錯誤: 尚未安裝 python-docx。請執行 pip install python-docx")
            return self.get_mock_data()

        # 2. 檢查檔案是否存在
        if not os.path.exists(filename):
            print(f"❌ 找不到檔案: '{filename}'")
            print("請確認 Word 檔是否放在同一個資料夾，且名稱完全正確。")
            return self.get_mock_data()

        print(f"📂 正在讀取檔案: {filename} ...")
        try:
            document = Document(filename)
        except Exception as e:
            print(f"❌ 讀取檔案失敗: {e}")
            return self.get_mock_data()

        processed_data = []
        current_day = 1
        word_count = 0
        table_count = len(document.tables)
        
        if table_count > 0:
            print(f"發現 {table_count} 個表格，正在解析...")
            
            for t_idx, table in enumerate(document.tables):
                # 每個表格前幾行稍微印出來 Debug
                # if t_idx == 0: print("正在檢查第一個表格結構...")

                for r_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    
                    # 過濾掉空行
                    if not any(cells): 
                        continue
                    
                    row_text = "".join(cells)
                    
                    # 偵測天數 (例如 "Day 1" 或 "第 1 天")
                    if "Day" in row_text or ("第" in row_text and "天" in row_text):
                        nums = re.findall(r'\d+', row_text)
                        if nums:
                            current_day = int(nums[0])
                            # print(f"--> 切換至第 {current_day} 天")
                        continue

                    # --- 強力解析邏輯 ---
                    word_cand = ""
                    ipa_cand = ""
                    mean_cand = ""
                    sent_cand = ""
                    
                    # 策略：逐格分析內容特性
                    for cell_text in cells:
                        if not cell_text: continue
                        
                        # 1. 如果包含中文 -> 很大機率是意思
                        if re.search(r'[\u4e00-\u9fff]', cell_text):
                            # 如果字數太多，可能是例句的中文翻譯，這裡簡單判斷長度
                            if len(cell_text) < 50:
                                if not mean_cand: mean_cand = cell_text
                            
                        # 2. 如果包含音標符號 / 或 [ -> 音標
                        elif ('/' in cell_text or '[' in cell_text) and len(cell_text) < 30:
                             if not ipa_cand: ipa_cand = cell_text

                        # 3. 如果是英文長句 (含空格) -> 例句
                        elif len(cell_text.split()) > 3:
                            if not sent_cand: sent_cand = cell_text
                            
                        # 4. 如果是英文短字 -> 可能是單字
                        # 允許包含一點雜訊(如數字)，稍後清理
                        elif re.search(r'[a-zA-Z]', cell_text):
                            # 排除太短的 (如編號 a, b) 除非是 a, I 等字
                            clean_text = self.clean_word_text(cell_text)
                            if len(clean_text) >= 1:
                                if not word_cand: word_cand = clean_text

                    # 只要有抓到單字，我們就收錄 (即使沒有意思或例句)
                    if word_cand:
                        # 排除標題行 (例如標題就是 "Word")
                        if word_cand.lower() in ['word', 'vocabulary', '單字']:
                            continue
                            
                        word_count += 1
                        # 如果是前幾筆，印出來讓用戶安心
                        if word_count <= 3:
                            print(f"   [範例] 抓到: {word_cand} ({mean_cand})")

                        entry = {
                            "id": word_count,
                            "day_number": current_day,
                            "word": word_cand,
                            "ipa": ipa_cand,
                            "meaning": mean_cand or "自訂", # 防呆
                            "sentence": sent_cand or f"Example for {word_cand}", # 防呆
                            "syllables": self.get_syllables(word_cand)
                        }
                        processed_data.append(entry)

        else:
            print("⚠️ 未發現表格，請確認 Word 檔是否使用表格排版。")

        if not processed_data:
            print("⚠️ 依然沒有抓到單字。")
            print("可能原因：表格格式太特殊。")
            return self.get_mock_data()

        return processed_data

    def get_mock_data(self):
        """生成範例資料 (備用)"""
        print("--> 生成 4 筆模擬資料...")
        return [
            {"id":1, "day_number":1, "word":"ability", "ipa":"/əˈbɪləti/", "meaning":"能力", "sentence":"She has the ability...", "syllables":["a","bil","i","ty"]},
            {"id":2, "day_number":1, "word":"abroad", "ipa":"/əˈbrɔːd/", "meaning":"在國外", "sentence":"Study abroad...", "syllables":["a","broad"]},
            {"id":3, "day_number":1, "word":"accept", "ipa":"/əkˈsept/", "meaning":"接受", "sentence":"Accept apology...", "syllables":["ac","cept"]},
            {"id":4, "day_number":2, "word":"accident", "ipa":"/ˈæksɪdənt/", "meaning":"意外", "sentence":"Car accident...", "syllables":["ac","ci","dent"]}
        ]

    def export_to_json(self, data, filename="pet_vocab_db.json"):
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"✅ 成功導出 {len(data)} 筆資料至 {filename}")

if __name__ == "__main__":
    processor = PetVocabProcessor()
    
    # 請確認這裡的檔名跟您桌面上的檔案一模一樣
    docx_filename = "更新版PET28天.docx" 
    
    final_data = processor.parse_docx(docx_filename)
    
    if final_data:
        processor.export_to_json(final_data)
        print("完成！請打開 pet_vocab_db.json 複製內容。")